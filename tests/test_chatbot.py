import pytest
import requests
import json
import logging
from docx import Document
from datetime import datetime


API_URL = "http://127.0.0.1:5000/api/v1/chat"  # Change this to your running endpoint
SESSION_URL = "http://127.0.0.1:5000/api/v1/session/create"  # Endpoint to get new session

doc = Document()
doc.add_heading("Chatbot Test Results", 0)
doc.add_paragraph(f"Generated on: {datetime.now()}\n")

# Load prompts from JSON file
with open("test_prompts_bravur.json", "r", encoding="utf-8") as f:
    test_cases = json.load(f)

def is_hallucination(response: str, expected_keywords: list) -> bool:
    """
    Basic heuristic to flag hallucinations: if expected keywords (or phrases) are missing.
    """
    response_lower = response.lower()
    return not any(keyword.lower() in response_lower for keyword in expected_keywords)

@pytest.fixture(scope="module")
def session_id():
    try:
        res = requests.post(SESSION_URL)  # or .get() depending on your API
        res.raise_for_status()
        session_id = res.json().get("session_id")

        # Change language to English for testing
        lang_res = requests.post("http://127.0.0.1:5000/api/v1/language_change", data={
            "session_id": session_id,
            "language": "en-US",
            "from_language": "nl-NL",
            "to_language": "en-US"
        })
        lang_res.raise_for_status()

        # Accept consent for this session
        consent_res = requests.post("http://127.0.0.1:5000/api/v1/consent/accept", json={"session_id": session_id})
        consent_res.raise_for_status()
        return session_id
    except Exception as e:
        logging.error(f"Failed to get session ID: {e}")
        pytest.skip("Skipping tests because no session ID could be obtained")
@pytest.mark.parametrize("test_case", test_cases)
def test_chatbot_responses(test_case, session_id):
    prompt = test_case["prompt"]
    expected_keywords = test_case.get("expected_keywords", [])
    expect_error = test_case.get("expect_error", False)

    try:
        res = requests.post(
            API_URL,
            headers={"Content-Type": "application/x-www-form-urlencoded"},
            data={"user_input": prompt, "session_id": session_id, "language": "en",}
        )

        assert res.status_code == 200, f"API Error: {res.status_code}"

        # ✅ Try parsing JSON, else fallback to plain text
        content_type = res.headers.get("Content-Type", "")
        if "application/json" in content_type:
            response_text = res.json().get("response", "")
        else:
            print(f"⚠️ Expected JSON but got: {content_type}")
            response_text = res.text.strip()

        if not response_text:
            raise ValueError("Received an empty response.")

        # ✅ Language verification (optional check)
        if not response_text.lower().startswith("bravur") and not response_text.isascii():
            raise AssertionError("Response is likely not in English.")

        # ✅ Validate response
        if expect_error:
            assert "[Error" in response_text or "couldn’t find" in response_text.lower()
        elif expected_keywords:
            assert not is_hallucination(response_text, expected_keywords), f"Possible hallucination in: {response_text}"
        else:
            assert response_text.strip(), "Empty response received"

        # Log passed test
        doc.add_heading(f"Prompt: {prompt}", level=2)
        doc.add_paragraph(f"Response:\n{response_text}")
        doc.add_paragraph("✅ Test passed.\n")


    except Exception as e:
        doc.add_heading(f"Prompt: {prompt}", level=2)
        doc.add_paragraph(f"Raw Response:\n{res.text}")
        doc.add_paragraph(f"❌ Test failed. Exception:\n{e}")

        print("⚠️ Failed during test execution.")
        print(f"Prompt: {prompt}")
        print(f"Raw response:\n{res.text}")
        print(f"Exception:\n{e}")
        logging.exception(f"Test case failed: {prompt}")
        assert False, f"Exception occurred during test: {e}"

@pytest.hookimpl(trylast=True)
def pytest_sessionfinish(session, exitstatus):
    doc.add_paragraph(f"\nTest session finished with status: {exitstatus}")
    doc.save("chatbot_test_results.docx")
